"""Document ingestion: extract -> chunk -> embed -> store (project_specs.md §10).

Robustness (constraint #8): one bad PDF page is logged and skipped; an
already-ingested file (same sha256) is skipped without re-embedding. Catastrophic
failures (embedding/DB) propagate to the handler, which replies to the user.
"""

from __future__ import annotations

import hashlib
import logging
import time
from dataclasses import dataclass
from io import BytesIO
from uuid import UUID

from docx import Document
from pypdf import PdfReader

from bot.models import Chunk
from bot.rag.chunker import TextChunk, chunk_text
from bot.services.embeddings import EmbeddingService
from bot.services.supabase_client import Database

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = ("pdf", "docx", "txt")


@dataclass
class IngestResult:
    filename: str
    chunks_added: int
    elapsed_seconds: float
    skipped: bool = False
    source_id: UUID | None = None


def _extract_segments(file_bytes: bytes, file_type: str) -> list[tuple[int | None, str]]:
    """Return ``[(page_no | None, text), ...]``; one bad PDF page is skipped."""
    if file_type == "txt":
        return [(None, file_bytes.decode("utf-8", errors="replace"))]
    if file_type == "docx":
        doc = Document(BytesIO(file_bytes))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text.strip())
        return [(None, "\n".join(parts))]
    if file_type == "pdf":
        segments: list[tuple[int | None, str]] = []
        reader = PdfReader(BytesIO(file_bytes))
        for page_no, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                logger.warning("PDF page %d failed to extract; skipping", page_no)
                continue
            if text.strip():
                segments.append((page_no, text))
        return segments
    raise ValueError(f"Unsupported file_type: {file_type}")


async def ingest_document(
    *,
    db: Database,
    embeddings: EmbeddingService,
    file_bytes: bytes,
    filename: str,
    file_type: str,
    uploaded_by: int,
    chunk_size_tokens: int,
    overlap_tokens: int,
    priority: int = 0,
) -> IngestResult:
    """Extract, chunk, embed, and atomically store one document."""
    started = time.perf_counter()
    sha256 = hashlib.sha256(file_bytes).hexdigest()

    existing = await db.find_active_source_by_hash(sha256)
    if existing is not None:
        logger.info("'%s' already ingested (sha match); skipping", filename)
        return IngestResult(
            filename,
            existing.chunk_count,
            time.perf_counter() - started,
            skipped=True,
            source_id=existing.id,
        )

    pieces: list[tuple[int | None, TextChunk]] = []
    for page_no, text in _extract_segments(file_bytes, file_type):
        for tc in chunk_text(
            text, chunk_size_tokens=chunk_size_tokens, overlap_tokens=overlap_tokens
        ):
            pieces.append((page_no, tc))

    if not pieces:
        logger.warning("No extractable text in '%s'", filename)
        return IngestResult(filename, 0, time.perf_counter() - started)

    vectors = await embeddings.embed_documents([tc.text for _, tc in pieces])

    chunks = [
        Chunk(
            chunk_index=i,
            content=tc.text,
            embedding=vector,
            token_count=tc.token_estimate,
            metadata={"page": page_no, "char_start": tc.char_start, "char_end": tc.char_end},
            priority=priority,
        )
        for i, ((page_no, tc), vector) in enumerate(zip(pieces, vectors, strict=True))
    ]

    source_id = await db.ingest_source_with_chunks(
        filename=filename,
        file_type=file_type,
        uploaded_by=uploaded_by,
        sha256=sha256,
        priority=priority,
        chunks=chunks,
    )
    logger.info("Ingested '%s': %d chunks", filename, len(chunks))
    return IngestResult(filename, len(chunks), time.perf_counter() - started, source_id=source_id)
